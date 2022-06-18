from docx import Document
# from docx.shared import Inches
import pandas as pd
import win32com.client
from docx.text.paragraph import Paragraph
import numpy as np


def make_client_invoice(employer, row_count, dataf):
    document = Document('C:/Users/lgiib/Downloads/KWANIA DEMAND LETTER.docx')
    # print(document.paragraphs)
    # document.add_picture('brand_logo.png', width=Inches(1))
    # document.add_heading('Invoice', 0)
    # p1 = document.add_paragraph('Dear ')
    # p1.add_run(name).bold = True
    # p1.add_run(',')
    p1 = document.paragraphs[2]
    r1 = p1.runs[0]
    r2 = p1.runs[1]
    r1.text = employer
    r2.text = ''

    # p2 = document.add_paragraph('Please find attached invoice for your recent purchase of ')
    # p2.add_run(str(unit)).bold = True
    # p2.add_run(' units of ')
    # p2.add_run(product).bold = True
    # p2.add_run('.')
    p2 = get_paragraph(document.paragraphs, 'Below is a summary table per period.')

    # [document.add_paragraph('') for _ in range(2)]

    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    # hdr_cells[0].text = 'EMPLOYER'
    hdr_cells[0].text = 'PERIOD'
    hdr_cells[1].text = 'EXPECTATION (UGX)'
    hdr_cells[2].text = 'RECEIVED (UGX)'
    hdr_cells[3].text = 'VARIANCE (UGX)'
    for x in range(4):
        hdr_cells[x].paragraphs[0].runs[0].font.bold = True

    t = int(row_count)
    for z in range(t):
        # print(f'z is {z}')
        # print(f't is {t}')
        # print(dataf.columns.values.tolist())
        row_cellsz = table.add_row().cells
        # row_cellsz[0].text = employer
        row_cellsz[0].text = f'{dataf.iloc[z]["PERIOD"]}'  # dataf.iloc[z, 4]
        row_cellsz[1].text = f'{"{:,}".format(int(dataf.iloc[z]["EXPECTATION"]))}'  # dataf.iloc[z, 1]
        row_cellsz[2].text = f'{"{:,}".format(int(dataf.iloc[z]["RECEIVED"]))}'  # dataf.iloc[z, 2]
        row_cellsz[3].text = f'{"{:,}".format(int(dataf.iloc[z]["VARIANCE"]))}'  # dataf.iloc[z, 3]
        if z+1 == t:
            row_cellsz = table.add_row().cells
            row_cellsz[0].text = 'TOTAL'
            row_cellsz[1].text = f'{"{:,}".format(int(dataf["EXPECTATION"].sum()))}'
            row_cellsz[2].text = f'{"{:,}".format(int(dataf["RECEIVED"].sum()))}'
            row_cellsz[3].text = f'{"{:,}".format(int(dataf["VARIANCE"].sum()))}'
            for q in range(4):
                row_cellsz[q].paragraphs[-1].runs[-1].font.bold = True


    # row_cells = table.add_row().cells
    # row_cells[0].text = 'TOTAL'
    # row_cells[1].text = f'{"{:,}".format(int(dataf["EXPECTATION"].sum()))}'
    # row_cells[2].text = f'{"{:,}".format(int(dataf["RECEIVED"].sum()))}'
    # row_cells[3].text = f'{"{:,}".format(int(dataf["VARIANCE"].sum()))}'
    # # https://stackoverflow.com/questions/41286569/get-total-of-pandas-column
    # for q in range(4):
    #     row_cells[q].paragraphs[-1].runs[-1].font.bold = True

    table.style = 'Table Grid'

    move_table_after(table, p2)

    # p3 = table_insert_paragraph_after(table)
    # [document.add_paragraph('') for _ in range(2)]

    # [document.add_paragraph('') for _ in range(2)]
    # [document.add_paragraph('') for _ in range(10)]
    #
    # document.add_paragraph('We appreciate your business and and please come again!')
    # document.add_paragraph('Sincerely')
    # document.add_paragraph('Jay')

    document.save(f'C:/Users/lgiib/Downloads/Demands/{employer}.docx')


# def docx_to_pdf(src, dst):
#     word = win32com.client.Dispatch("Word.Application")
#     wdFormatPDF = 17
#     doc = word.Documents.Open(src)
#     doc.SaveAs(dst, FileFormat=wdFormatPDF)
#     doc.Close()
#     word.Quit()


# def send_email(name, to_addr, attachment):
#     outlook = win32com.client.Dispatch("Outlook.Application")
#     mail = outlook.CreateItem(0)
#     mail.To = to_addr  # 'amznbotnotification@gmail.com'
#     mail.Subject = 'Invoice from PythonInOffice'
#     mail.Body = f'Dear {name}, Please find attached invoice'
#     mail.Attachments.Add(attachment)
#     mail.Send()

def get_paragraph(paras, text):
    """Return the paragraph where the text resides

        Args:
            paras(document.paragraphs): All the paragraphs in the document
            text (str): The text in the paragraph to match

        """
    for para in paras:
        if text in para.text:
            return para
    raise KeyError("The text cannot be found anywhere in the document")


def move_table_after(tableii, paragraphii):
    tbl, p = tableii._tbl, paragraphii._p
    p.addnext(tbl)


# however you get this paragraph
# tableII = document.add_table(...)
# move_table_after(tableII, paragraph)
# source: https://github.com/python-openxml/python-docx/issues/156

def table_insert_paragraph_after(table):
    """Return new `Paragraph` object inserted directly after `table`.

    `table` must already be immediately followed by a paragraph. So
    This won't work for a table followed by another table or a table
    at the end of the document.
    """
    p = table._tbl.getnext()
    paragraph = Paragraph(p, table._parent)
    return paragraph.insert_paragraph_before()


# paragraph = table_insert_paragraph_after(table)


df = pd.read_excel(r'C:/Users/lgiib/Downloads/EVP.xlsx')

vote_data = df["EMPLOYER"].unique()  # Brings out unique values

for i in vote_data:
    # print(i)
    a = df[df["EMPLOYER"].str.contains(i)]  # Filter for all that contain current i into a df
    # a.replace(r'^\s*$', np.nan, regex=True) # Replace empty values with NaN
    # b = a.fillna(0)  # Replace NaN with 0
    a['RECEIVED'] = a['RECEIVED'].fillna(0)
    # https://datatofish.com/replace-nan-values-with-zeros/
    # print(a)
    result = a["EMPLOYER"].value_counts()  # Returns a series in a column
    result_two = a.value_counts()  # Returns a series in a table, however deletes rows with NaN
    # print(result)
    # print(type(result))
    r = result_two.to_frame()
    # r.drop(r.columns[len(r.columns) - 1], axis=1, inplace=True)
    # w = r.replace(r'^\s*$', np.nan, regex=True)
    # print(w)
    r.to_csv(r'C:/Users/lgiib/Downloads/EVP.csv')  # These two help eliminate time on the period column
    s = pd.read_csv(r'C:/Users/lgiib/Downloads/EVP.csv')
    # print(s)
    # print(r.iloc[1:])
    # print(result[0])
    # print(i)
    # https://datascienceparichay.com/article/convert-pandas-series-to-a-dictionary/
    result_dict = result.to_dict()
    # print(result_dict)
    try:
        make_client_invoice(i, result_dict[i], s)
    except KeyError:
        print(i)
